import os
import re
import sys
import shutil
from pathlib import Path
from typing import Iterable, Tuple

# =====================
# CONFIG — KNOBS
# =====================
ROOT_DIR = r"D:\ChangeMe\root"          # root folder to process
TARGET_WORD = "QWEN"                      # base target (case-insensitive match of this exact sequence)
REPLACEMENT_WORD = "KIMI"                 # base replacement; script preserves the case pattern of each match
MAX_DEPTH = -1                             # 0 = only ROOT_DIR; 5 = up to 5 levels; -1 = unlimited

# File types to include (treated as text unless specialized handlers exist)
INCLUDE_EXTS = {
    ".md", ".txt", ".csv", ".json", ".yaml", ".yml", ".ini", ".cfg", ".xml", ".html",
    ".docx",  # handled via python-docx
    ".xlsx",  # handled via openpyxl
}

EXCLUDE_EXTS = {
    ".rtf", ".doc", ".xls"  # legacy / risky formats excluded per requirements
}

DRY_RUN = True                            # True = preview only; False = actually write changes
BACKUP_BEFORE_WRITE = False               # Optional: copy changed files to a backup folder
BACKUP_DIR = r"D:\ChangeMe\backup"      # used only if BACKUP_BEFORE_WRITE is True

# =====================
# Utilities
# =====================
UPPER = str.upper
LOWER = str.lower

_case_insensitive_pat = None  # compiled at runtime


def compile_pattern(target: str) -> re.Pattern:
    """Compile a case-insensitive regex that matches the exact target sequence (as substring)."""
    # Escape target to treat it literally
    escaped = re.escape(target)
    return re.compile(escaped, re.IGNORECASE)


def apply_case_pattern(sample: str, base: str) -> str:
    """Return base transformed to mimic the per-character case pattern of sample.

    Rules:
    - For i < len(sample): if sample[i] is upper -> base[i].upper(); if lower -> base[i].lower(); else keep base[i].
    - For i >= len(sample): repeat last known case pattern; default to lower if unknown.
    - If base shorter than sample, we still map per available chars.
    """
    out_chars = []
    last_is_upper = None

    for i, ch in enumerate(base):
        if i < len(sample):
            s = sample[i]
            if s.isalpha():
                out_chars.append(ch.upper() if s.isupper() else ch.lower())
                last_is_upper = s.isupper()
            else:
                out_chars.append(ch)
            continue
        # Beyond the sample length: extend using the last observed case preference
        if last_is_upper is True:
            out_chars.append(ch.upper())
        elif last_is_upper is False:
            out_chars.append(ch.lower())
        else:
            out_chars.append(ch)  # no guidance; keep as-is
    return "".join(out_chars)


def replace_preserving_case(text: str, pat: re.Pattern, replacement_base: str) -> Tuple[str, int]:
    """Replace all occurrences matched by pat with replacement_base while preserving each match's case pattern.
    Returns new_text, count.
    """
    def _sub(m: re.Match) -> str:
        sample = m.group(0)
        return apply_case_pattern(sample, replacement_base)

    return pat.subn(_sub, text)


def within_depth(root: Path, path: Path, max_depth: int) -> bool:
    if max_depth < 0:
        return True
    try:
        rel = path.relative_to(root)
    except ValueError:
        return False
    # Depth = number of parts in rel.parent (directories under root)
    if rel == Path("."):
        depth = 0
    else:
        depth = len(rel.parts) - 1 if path.is_file() else len(rel.parts)
    return depth <= max_depth


# =====================
# Text handlers
# =====================

def process_text_file(p: Path, pat: re.Pattern) -> Tuple[bool, int]:
    """Process generic text-like files (utf-8 with fallback). Returns (changed, replacements)."""
    try:
        data = p.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # fallback as latin-1 to avoid crash; user can re-save if needed
        data = p.read_text(encoding="latin-1")
    new_data, n = replace_preserving_case(data, pat, REPLACEMENT_WORD)
    if n > 0 and not DRY_RUN:
        if BACKUP_BEFORE_WRITE:
            backup_file(p)
        p.write_text(new_data, encoding="utf-8")
    return (n > 0, n)


# =====================
# DOCX handler (python-docx)
# =====================

def process_docx_file(p: Path, pat: re.Pattern) -> Tuple[bool, int]:
    try:
        from docx import Document
    except Exception as e:
        print(f"[WARN] python-docx not available for {p.name}: {e}")
        return (False, 0)

    doc = Document(str(p))
    total = 0

    def replace_in_text_obj(text_obj) -> int:
        # Replace inside a single string and push back
        new_text, n = replace_preserving_case(text_obj.text, pat, REPLACEMENT_WORD)
        if n:
            text_obj.text = new_text
        return n

    # Paragraphs
    for para in doc.paragraphs:
        total += replace_in_text_obj(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += replace_in_text_obj(para)

    if total > 0 and not DRY_RUN:
        if BACKUP_BEFORE_WRITE:
            backup_file(p)
        doc.save(str(p))
    return (total > 0, total)


# =====================
# XLSX handler (openpyxl)
# =====================

def process_xlsx_file(p: Path, pat: re.Pattern) -> Tuple[bool, int]:
    try:
        import openpyxl
    except Exception as e:
        print(f"[WARN] openpyxl not available for {p.name}: {e}")
        return (False, 0)

    wb = openpyxl.load_workbook(filename=str(p))
    total = 0

    for ws in wb.worksheets:
        # Sheet title
        new_title, n = replace_preserving_case(ws.title, pat, REPLACEMENT_WORD)
        if n:
            try:
                ws.title = new_title
                total += n
            except Exception:
                pass  # sheet title collisions/limits → ignore

        # Cell values
        for row in ws.iter_rows(values_only=False):
            for cell in row:
                if isinstance(cell.value, str):
                    new_val, n = replace_preserving_case(cell.value, pat, REPLACEMENT_WORD)
                    if n:
                        cell.value = new_val
                        total += n

    if total > 0 and not DRY_RUN:
        if BACKUP_BEFORE_WRITE:
            backup_file(p)
        wb.save(str(p))
    return (total > 0, total)


# =====================
# Filename renaming (case-preserving on base name only)
# =====================

def rename_file_if_needed(p: Path, pat: re.Pattern) -> Tuple[bool, Path]:
    name = p.name
    new_name, n = replace_preserving_case(name, pat, REPLACEMENT_WORD)
    if n == 0:
        return (False, p)

    new_path = p.with_name(resolve_collision(p.parent, new_name))
    if not DRY_RUN:
        try:
            p.rename(new_path)
        except OSError:
            # cross-device moves or locked files → fallback copy+remove
            temp = new_path
            shutil.copy2(str(p), str(temp))
            p.unlink(missing_ok=True)
    return (True, new_path)


def resolve_collision(folder: Path, desired_name: str) -> str:
    base, ext = os.path.splitext(desired_name)
    candidate = desired_name
    i = 1
    while (folder / candidate).exists():
        candidate = f"{base}_{i}{ext}"
        i += 1
    return candidate


def backup_file(p: Path) -> None:
    if not BACKUP_BEFORE_WRITE:
        return
    dst_root = Path(BACKUP_DIR)
    rel = p.relative_to(Path(ROOT_DIR))
    dst = dst_root / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(p), str(dst))


# =====================
# Main driver
# =====================

def process_file(p: Path, pat: re.Pattern) -> Tuple[int, bool, Path]:
    """Return (replacements_count, renamed, final_path)."""
    ext = p.suffix.lower()
    changed = 0

    if ext in {".docx"}:
        _, changed = process_docx_file(p, pat)
    elif ext in {".xlsx"}:
        _, changed = process_xlsx_file(p, pat)
    elif ext in INCLUDE_EXTS and ext not in {".docx", ".xlsx"}:
        _, changed = process_text_file(p, pat)
    else:
        return (0, False, p)

    renamed, new_path = rename_file_if_needed(p, pat)
    return (changed, renamed, new_path)


def walk_paths(root: Path, max_depth: int) -> Iterable[Path]:
    for dirpath, dirnames, filenames in os.walk(root):
        dir_path = Path(dirpath)
        # Enforce depth limit: filter traversal by pruning dirnames in-place
        if max_depth >= 0:
            depth = len(dir_path.relative_to(root).parts) if dir_path != root else 0
            if depth >= max_depth:
                dirnames[:] = []  # don't descend further
        for fn in filenames:
            yield dir_path / fn


def main():
    global _case_insensitive_pat
    root = Path(ROOT_DIR)
    assert root.exists(), f"Root folder does not exist: {ROOT_DIR}"

    _case_insensitive_pat = compile_pattern(TARGET_WORD)

    if BACKUP_BEFORE_WRITE:
        Path(BACKUP_DIR).mkdir(parents=True, exist_ok=True)

    total_seen = 0
    total_changed = 0
    total_renamed = 0

    for p in walk_paths(root, MAX_DEPTH):
        if p.suffix.lower() in EXCLUDE_EXTS:
            continue
        if p.suffix.lower() not in INCLUDE_EXTS:
            continue
        if not within_depth(root, p, MAX_DEPTH):
            continue

        total_seen += 1
        changed, renamed, _ = process_file(p, _case_insensitive_pat)
        total_changed += changed
        total_renamed += int(renamed)

        if DRY_RUN and (changed or renamed):
            print(f"[DRY] {p}")
            if changed:
                print(f"   - content replacements: {changed}")
            if renamed:
                print(f"   - will rename filename")

    print("\n===== SUMMARY =====")
    print(f"Files scanned        : {total_seen}")
    print(f"Content replacements : {total_changed}")
    print(f"Files renamed        : {total_renamed}")
    print(f"Depth limit          : {MAX_DEPTH}")
    print(f"Dry run              : {DRY_RUN}")


if __name__ == "__main__":
    main()
